from __future__ import annotations

import re
import sys
from pathlib import Path
from urllib.parse import unquote

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_paragraph_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "Normal"
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def apply_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(22)
    title.font.bold = True

    for idx, size in [(1, 16), (2, 14), (3, 12)]:
        style = document.styles[f"Heading {idx}"]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True


def add_inline(paragraph, node, bold=False, italic=False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        return

    if not isinstance(node, Tag):
        return

    next_bold = bold or node.name in {"strong", "b"}
    next_italic = italic or node.name in {"em", "i"}

    if node.name == "code":
        run = paragraph.add_run(node.get_text())
        run.font.name = "Consolas"
        run.bold = next_bold
        run.italic = next_italic
        return

    if node.name == "br":
        paragraph.add_run("\n")
        return

    if node.name == "a":
        text = node.get_text(strip=False)
        href = node.get("href", "")
        combined = f"{text} ({href})" if href and href not in text else text
        run = paragraph.add_run(combined)
        run.bold = next_bold
        run.italic = next_italic
        return

    for child in node.children:
        add_inline(paragraph, child, next_bold, next_italic)


def add_image(document: Document, src: str, base_dir: Path, max_width: float = 6.7) -> None:
    original_path = base_dir / Path(unquote(src))
    labelled_path = base_dir / "labelled" / Path(unquote(src))
    image_path = labelled_path if labelled_path.exists() else original_path
    if not image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f"[Missing image: {original_path.name}]").italic = True
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(max_width))


def handle_paragraph(document: Document, element: Tag, base_dir: Path) -> None:
    images = element.find_all("img", recursive=False)
    if images and not element.get_text(strip=True):
        for image in images:
            add_image(document, image.get("src", ""), base_dir)
        return

    paragraph = document.add_paragraph()
    paragraph.style = "Normal"
    for child in element.children:
        add_inline(paragraph, child)

    text = paragraph.text.strip()
    if not text:
        paragraph._element.getparent().remove(paragraph._element)
        return

    if re.match(r"^\*\*Student (name|ID):", text, flags=re.IGNORECASE):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text.startswith("Figure "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def handle_list(document: Document, element: Tag) -> None:
    style = "List Number" if element.name == "ol" else "List Bullet"
    for li in element.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style)
        for child in li.children:
            add_inline(paragraph, child)


def handle_table(document: Document, element: Tag) -> None:
    rows = element.find_all("tr")
    if not rows:
        return

    col_count = max(len(row.find_all(["th", "td"])) for row in rows)
    table = document.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        doc_row = table.add_row()
        for col_idx, cell in enumerate(cells):
            set_cell_text(doc_row.cells[col_idx], cell.get_text(" ", strip=True))
            if row_idx == 0 or cell.name == "th":
                for run in doc_row.cells[col_idx].paragraphs[0].runs:
                    run.bold = True


def convert(markdown_path: Path, output_path: Path) -> None:
    text = markdown_path.read_text(encoding="utf-8")
    html = markdown.markdown(text, extensions=["tables", "sane_lists"])
    soup = BeautifulSoup(html, "html.parser")

    document = Document()
    apply_styles(document)
    document.core_properties.title = "CMP-X305 Cyber Security Coursework Portfolio 1"

    title_seen = False
    for element in soup.children:
        if not isinstance(element, Tag):
            continue

        if element.name == "h1":
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(element.get_text(strip=True))
            title_seen = True
            continue

        if element.name in {"h2", "h3", "h4"}:
            level = {"h2": 1, "h3": 2, "h4": 3}[element.name]
            paragraph = document.add_heading(element.get_text(strip=True), level=level)
            set_paragraph_keep_with_next(paragraph)
            if title_seen and level == 1 and document.paragraphs[-2].style.name != "Title":
                pass
            continue

        if element.name == "p":
            handle_paragraph(document, element, markdown_path.parent)
            continue

        if element.name in {"ul", "ol"}:
            handle_list(document, element)
            continue

        if element.name == "table":
            handle_table(document, element)
            continue

        if element.name == "hr":
            document.add_paragraph()
            continue

    footer = document.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("Page ")
    add_page_number(footer_p)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    markdown_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("CMP-X305_coursework_report.md")
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("CMP-X305_coursework_report.docx")
    convert(markdown_path, output_path)
    print(output_path.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
